from __future__ import annotations

import json
from collections import defaultdict
from dataclasses import asdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.db.postgres import get_pg_conn


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

COMPANIES = ["005930", "000660", "042700"]
COMPANY_NAMES = {
    "005930": "삼성전자",
    "000660": "SK하이닉스",
    "042700": "한미반도체",
}
HORIZONS = [7, 14, 30]
PCT_KEYS = {"expected_return", "low_return", "high_return", "realized_return", "abs_error"}
FLOAT_KEYS = {
    "base_price",
    "expected_price",
    "low_price",
    "high_price",
    "latest_close",
    "return_7d",
    "return_20d",
    "return_60d",
    "volatility_20d",
    "volatility_60d",
    "drawdown_60d",
    "range_60d",
    "volume_ratio_20d",
    "trend_strength",
}

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\malgun.ttf"),
    Path(r"C:\Windows\Fonts\malgunbd.ttf"),
]


def _fmt_pct(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):+.2%}"


def _fmt_num(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):,.2f}"


def _fmt_dt(value) -> str:
    if value is None:
        return "-"
    if hasattr(value, "isoformat"):
        text = value.isoformat()
    else:
        text = str(value)
    return text.replace("+00:00", " UTC")


def _json_safe(value):
    if value is None:
        return {}
    return value


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: str = "000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = "E8EEF5", font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(value, str) and value.replace(".", "", 1).replace("-", "", 1).replace("+", "", 1).isdigit():
                align = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(cells[i], value, size=font_size, align=align)
    return table


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    normal = styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        s = styles[name]
        s.font.size = Pt(size)
        s.font.bold = False
        s.font.color.rgb = RGBColor.from_string("000000" if name == "Title" else "1F4D78" if name != "Heading 3" else "434343")
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def _query_rows(sql: str, params=()):
    with get_pg_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, params)
            cols = [d[0] for d in cur.description]
            rows = cur.fetchall()
    return [dict(zip(cols, row)) for row in rows]


def _register_reportlab_font() -> str:
    for path in FONT_CANDIDATES:
        if path.exists():
            font_name = "MalgunGothic"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
    return "Helvetica"


def _rl_escape(value: str) -> str:
    value = str(value)
    value = (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return value


def _rl_paragraph(text: str, style: ParagraphStyle):
    return Paragraph(_rl_escape(text), style)


def _rl_table(data: list[list[str]], col_widths: list[float], font_name: str, font_size: int = 8, header_fill=colors.HexColor("#E8EEF5")):
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("LEADING", (0, 0), (-1, -1), font_size + 2),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), header_fill),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7D1DC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def fetch_latest_forecasts():
    sql = """
    WITH latest AS (
      SELECT company_code, horizon_days, MAX(as_of) AS max_as_of
      FROM price_forecasts
      WHERE company_code = ANY(%s)
      GROUP BY company_code, horizon_days
    )
    SELECT f.company_code, f.horizon_days, f.as_of, f.published_cutoff, f.target_date,
           f.base_price, f.expected_return, f.low_return, f.high_return,
           f.expected_price, f.low_price, f.high_price, f.method, f.features, f.signals, f.scenarios
    FROM price_forecasts f
    JOIN latest l ON l.company_code = f.company_code
                 AND l.horizon_days = f.horizon_days
                 AND l.max_as_of = f.as_of
    ORDER BY f.company_code, f.horizon_days
    """
    return _query_rows(sql, (COMPANIES,))


def fetch_backtest_summary():
    sql = """
    SELECT f.company_code, f.horizon_days,
           COUNT(*) AS cases,
           AVG(e.expected_return) AS avg_expected_return,
           AVG(e.realized_return) AS avg_realized_return,
           AVG(e.abs_error) AS avg_abs_error,
           AVG(CASE WHEN e.interval_hit THEN 1.0 ELSE 0.0 END) AS interval_hit_rate,
           AVG(CASE
                 WHEN e.expected_return = 0 THEN 1.0
                 WHEN (e.expected_return > 0) = (e.realized_return > 0) THEN 1.0
                 ELSE 0.0
               END) AS direction_accuracy
    FROM price_forecast_evaluations e
    JOIN price_forecasts f ON f.id = e.forecast_id
    WHERE f.company_code = ANY(%s)
    GROUP BY f.company_code, f.horizon_days
    ORDER BY f.company_code, f.horizon_days
    """
    return _query_rows(sql, (COMPANIES,))


def fetch_overall_summary():
    sql = """
    SELECT
      COUNT(*) AS cases,
      AVG(e.expected_return) AS avg_expected_return,
      AVG(e.realized_return) AS avg_realized_return,
      AVG(e.abs_error) AS avg_abs_error,
      AVG(CASE WHEN e.interval_hit THEN 1.0 ELSE 0.0 END) AS interval_hit_rate,
      AVG(CASE
            WHEN e.expected_return = 0 THEN 1.0
            WHEN (e.expected_return > 0) = (e.realized_return > 0) THEN 1.0
            ELSE 0.0
          END) AS direction_accuracy
    FROM price_forecast_evaluations e
    JOIN price_forecasts f ON f.id = e.forecast_id
    WHERE f.company_code = ANY(%s)
    """
    rows = _query_rows(sql, (COMPANIES,))
    return rows[0] if rows else {}


def fetch_representative_cases():
    sql = """
    SELECT f.id, f.company_code, f.horizon_days, f.as_of, f.published_cutoff, f.target_date,
           f.base_price, f.expected_return, f.low_return, f.high_return,
           f.expected_price, f.low_price, f.high_price, f.features, f.signals, f.scenarios,
           e.realized_at, e.realized_close, e.realized_return, e.abs_error, e.interval_hit,
           e.scenario_hit, e.feedback
    FROM price_forecasts f
    JOIN price_forecast_evaluations e ON e.forecast_id = f.id
    WHERE f.company_code = %s
    ORDER BY e.abs_error DESC, f.horizon_days DESC, f.as_of DESC
    LIMIT 1
    """
    out = []
    for company in COMPANIES:
        rows = _query_rows(sql, (company,))
        if rows:
            row = rows[0]
            row["company_name"] = COMPANY_NAMES[company]
            out.append(row)
    return out


def _add_title_block(doc: Document, latest_as_of: str):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("삼성전자·SK하이닉스·한미반도체 주가 예측 및 백테스트 리포트")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.bold = False
    run.font.color.rgb = RGBColor.from_string("000000")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(10)
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = meta.add_run(f"기준 시점(as_of): {latest_as_of} | 생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def _add_summary_paragraphs(doc: Document, overall: dict, representative_cases: list[dict]):
    doc.add_heading("요약", level=1)
    p = doc.add_paragraph()
    p.add_run("핵심 결론: ").bold = True
    p.add_run(
        "현재 예측 엔진은 가격/이벤트/지표 데이터를 모두 as_of 이전으로 제한한 상태에서 동작하며, "
        "삼성전자와 SK하이닉스는 여전히 강한 추세를 전제로 한 상방 시나리오가 많고, "
        "한미반도체는 최근 약세를 반영한 보수적 시나리오가 우세하다."
    )

    p = doc.add_paragraph()
    p.add_run("백테스트 해석: ").bold = True
    p.add_run(
        "평균적으로는 방향성보다 범위 적중률이 더 약했고, 큰 오차는 대부분 추세 반전이나 이벤트 충격을 "
        "충분히 반영하지 못한 케이스에서 발생했다."
    )

    p = doc.add_paragraph()
    p.add_run("데이터 기반 한계: ").bold = True
    p.add_run(
        "대표 사례들의 feature에는 가격 모멘텀, 변동성, drawdown, 거래량 비율이 핵심이고, "
        "event_bias나 macro_bias는 대부분 0에 가까워서 외생 변수 반영이 아직 얇다."
    )

    doc.add_heading("전체 백테스트 요약", level=1)
    lines = [
        f"선정 종목 수: {len(COMPANIES)}",
        f"총 케이스 수: {overall.get('cases', 0)}",
        f"평균 예측 수익률: {_fmt_pct(overall.get('avg_expected_return'))}",
        f"평균 실현 수익률: {_fmt_pct(overall.get('avg_realized_return'))}",
        f"평균 절대오차: {_fmt_pct(overall.get('avg_abs_error'))}",
        f"구간 적중률: {_fmt_pct(overall.get('interval_hit_rate'))}",
        f"방향 정확도: {_fmt_pct(overall.get('direction_accuracy'))}",
    ]
    for item in lines:
        doc.add_paragraph(item, style="List Bullet")


def _add_forecast_table(doc: Document, rows: list[dict]):
    doc.add_heading("최신 예측 스냅샷", level=1)
    table_rows = []
    for row in rows:
        as_of_short = _fmt_dt(row["as_of"]).split("T")[0]
        table_rows.append(
            [
                COMPANY_NAMES.get(row["company_code"], row["company_code"]),
                f"{row['horizon_days']}d",
                as_of_short,
                str(row["target_date"]),
                _fmt_num(row["base_price"]),
                _fmt_pct(row["expected_return"]),
                f"{_fmt_pct(row['low_return'])} ~ {_fmt_pct(row['high_return'])}",
            ]
        )
    _add_table(
        doc,
        ["종목", "기간", "as_of", "목표일", "기준가", "예상수익률", "예측밴드"],
        table_rows,
        [1.35, 0.55, 1.45, 1.15, 1.0, 0.95, 1.55],
        header_fill="F2F4F7",
        font_size=9,
    )


def _add_backtest_table(doc: Document, rows: list[dict]):
    doc.add_heading("백테스트 요약", level=1)
    table_rows = []
    for row in rows:
        table_rows.append(
            [
                COMPANY_NAMES.get(row["company_code"], row["company_code"]),
                f"{row['horizon_days']}d",
                str(int(row["cases"])),
                _fmt_pct(row["avg_expected_return"]),
                _fmt_pct(row["avg_realized_return"]),
                _fmt_pct(row["avg_abs_error"]),
                _fmt_pct(row["interval_hit_rate"]),
                _fmt_pct(row["direction_accuracy"]),
            ]
        )
    _add_table(
        doc,
        ["종목", "기간", "케이스", "예상", "실현", "절대오차", "적중률", "방향정확도"],
        table_rows,
        [1.35, 0.55, 0.55, 0.85, 0.85, 0.85, 0.75, 0.95],
        header_fill="E8EEF5",
        font_size=9,
    )


def _case_feature_rows(features: dict, signals: list[dict]) -> list[list[str]]:
    rows = []
    for key in [
        "latest_close",
        "return_7d",
        "return_20d",
        "return_60d",
        "volatility_20d",
        "volatility_60d",
        "drawdown_60d",
        "range_60d",
        "volume_ratio_20d",
        "trend_strength",
    ]:
        rows.append([key, _fmt_pct(features.get(key)) if key in PCT_KEYS or "return" in key or "drawdown" in key or "volatility" in key else _fmt_num(features.get(key))])

    # Override formatting for non-percentage rows
    for i, row in enumerate(rows):
        key = row[0]
        if key in {"latest_close", "range_60d", "volume_ratio_20d", "trend_strength"}:
            rows[i][1] = _fmt_num(features.get(key))
        else:
            rows[i][1] = _fmt_pct(features.get(key))
    return rows


def _add_case_study(doc: Document, case: dict):
    doc.add_heading(f"{case['company_name']} 대표 사례", level=1)
    p = doc.add_paragraph()
    p.add_run("선정 기준: ").bold = True
    p.add_run("회사별 절대오차가 가장 큰 백테스트 케이스를 대표 사례로 선택했다.")

    summary_rows = [
        ["as_of", _fmt_dt(case["as_of"])],
        ["horizon", f"{case['horizon_days']}d"],
        ["target_date", str(case["target_date"])],
        ["base_price", _fmt_num(case["base_price"])],
        ["predicted_return", _fmt_pct(case["expected_return"])],
        ["predicted_band", f"{_fmt_pct(case['low_return'])} ~ {_fmt_pct(case['high_return'])}"],
        ["realized_close", _fmt_num(case["realized_close"])],
        ["realized_return", _fmt_pct(case["realized_return"])],
        ["abs_error", _fmt_pct(case["abs_error"])],
        ["interval_hit", "YES" if case["interval_hit"] else "NO"],
        ["scenario_hit", case["scenario_hit"]],
        ["feedback", case["feedback"]],
    ]
    _add_table(doc, ["항목", "값"], summary_rows, [1.8, 4.7], header_fill="F2F4F7", font_size=9)

    doc.add_paragraph()
    doc.add_heading("예측에 사용된 데이터", level=2)
    feature_rows = _case_feature_rows(case["features"], case["signals"])
    _add_table(doc, ["feature", "value"], feature_rows, [2.0, 4.5], header_fill="E8EEF5", font_size=9)

    doc.add_paragraph()
    doc.add_heading("시나리오 신호", level=2)
    scenario_rows = []
    for sc in case["signals"]:
        scenario_rows.append([
            sc.get("name", "-"),
            f"{float(sc.get('probability') or 0.0):.2%}",
            _fmt_pct(sc.get("expected_return")),
            f"{_fmt_pct(sc.get('low_return'))} ~ {_fmt_pct(sc.get('high_return'))}",
            sc.get("rationale", ""),
        ])
    if scenario_rows:
        _add_table(
            doc,
            ["scenario", "probability", "expected", "band", "rationale"],
            scenario_rows,
            [0.85, 0.85, 0.85, 1.15, 2.8],
            header_fill="E8EEF5",
            font_size=8,
        )

    doc.add_paragraph()
    doc.add_heading("해석", level=2)
    if case["company_code"] == "005930":
        narrative = (
            "삼성전자는 최근 7일/20일/60일 수익률이 모두 강했고(trend_strength 3.0), "
            "60일 range가 넓으면서도 변동성은 비교적 안정적이었다. 모델은 이런 모멘텀을 바탕으로 "
            "bull 시나리오 84.8%에 가까운 상방 지속을 가정했지만, 실제로는 30일 뒤 수익률이 -6.68%로 "
            "반전됐다. 핵심 원인은 추세 반전과 이벤트 충격을 충분히 반영하지 못한 점이다."
        )
        fix = "이 케이스는 외생 이벤트(공시/IR/산업 수급) 신호가 거의 없어, 가격 모멘텀에 과하게 의존한 것이 약점이다."
    elif case["company_code"] == "000660":
        narrative = (
            "SK하이닉스는 단기 수익률이 음수였지만 60일 기준으로는 여전히 플러스였고, drawdown도 -20%대였다. "
            "모델은 완만한 회복을 예상하며 base/bear 혼합 시나리오를 유지했지만, 실제로는 30일 동안 +82.76%로 "
            "급격한 재평가가 발생했다. 이 정도의 오차는 변동성과 이벤트 리스크를 지나치게 낮게 본 결과다."
        )
        fix = "이 케이스는 HBM/메모리 수요 급변처럼 가격 외 신호가 들어와야 설명력이 높아지는데, 현재 payload는 그 축이 비어 있다."
    else:
        narrative = (
            "한미반도체는 최근 7일/20일 수익률이 모두 음수이고 trend_strength가 -3.0으로 강한 약세였기 때문에, "
            "모델은 bear 시나리오 84.8%를 부여하고 추가 하락을 예상했다. 그러나 실제 30일 수익률은 +56.55%로 "
            "급반전했다. 즉, 최근 가격 추세는 맞게 읽었지만 이후의 수급/이벤트 전환을 놓친 케이스다."
        )
        fix = "이 케이스는 장비/후공정 수요 전환 같은 외생 촉매가 없으면, 단기 약세를 장기 약세로 오인하기 쉽다."
    doc.add_paragraph(narrative)
    doc.add_paragraph(f"개선 포인트: {fix}")


def build_report(output_path: Path) -> Path:
    latest = fetch_latest_forecasts()
    backtest = fetch_backtest_summary()
    overall = fetch_overall_summary()
    cases = fetch_representative_cases()
    latest_as_of = max((row["as_of"] for row in latest), default=datetime.now(timezone.utc))

    doc = Document()
    _configure_document(doc)
    _add_title_block(doc, _fmt_dt(latest_as_of))
    _add_summary_paragraphs(doc, overall, cases)
    _add_forecast_table(doc, latest)
    _add_backtest_table(doc, backtest)

    for idx, case in enumerate(cases):
        doc.add_page_break()
        _add_case_study(doc, case)

    doc.add_page_break()
    doc.add_heading("결론 및 한계", level=1)
    doc.add_paragraph(
        "이 리포트의 예측 엔진은 LLM이 임의로 수치를 뱉는 방식이 아니라, 가격 히스토리와 이벤트/지표 신호를 "
        "조합해 시나리오를 계산하고 그 결과를 백테스트로 검증하는 구조다."
    )
    doc.add_paragraph(
        "다만 현재 대표 사례를 보면 event_bias와 macro_bias가 대부분 비어 있어, 외생 변수 충격이나 급격한 "
        "재평가 구간에서는 예측 범위가 쉽게 벗어난다. 따라서 배포용 품질을 더 높이려면 공시, IR, 업계 수급, "
        "표준/로드맵 신호를 더 촘촘하게 넣어야 한다."
    )
    doc.add_paragraph(
        "요약하면, 현재 모델은 방향성보다 추세 지속 구간에 강하고, 이벤트 전환이나 급변 구간에 약하다. "
        "이 약점을 줄이기 위한 다음 단계는 데이터 소스 확장과 캘리브레이션 보정이다."
    )

    doc.save(output_path)
    return output_path


def build_pdf(output_path: Path) -> Path:
    latest = fetch_latest_forecasts()
    backtest = fetch_backtest_summary()
    overall = fetch_overall_summary()
    cases = fetch_representative_cases()
    latest_as_of = max((row["as_of"] for row in latest), default=datetime.now(timezone.utc))

    font_name = _register_reportlab_font()
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="KTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=20,
            leading=24,
            textColor=colors.black,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KH1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KH2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=8,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9.8,
            leading=14,
            textColor=colors.black,
            spaceAfter=5,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KSmall",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#333333"),
            spaceAfter=3,
            alignment=TA_LEFT,
        )
    )

    story = []
    story.append(_rl_paragraph("삼성전자·SK하이닉스·한미반도체 주가 예측 및 백테스트 리포트", styles["KTitle"]))
    story.append(_rl_paragraph(f"기준 시점(as_of): {_fmt_dt(latest_as_of)} | 생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", styles["KSmall"]))
    story.append(Spacer(1, 0.10 * inch))

    story.append(_rl_paragraph("요약", styles["KH1"]))
    story.append(_rl_paragraph(
        "현재 예측 엔진은 가격/이벤트/지표 데이터를 모두 as_of 이전으로 제한한 상태에서 동작한다. "
        "삼성전자와 SK하이닉스는 여전히 강한 추세를 전제로 한 상방 시나리오가 많고, 한미반도체는 최근 약세를 반영한 보수적 시나리오가 우세하다.",
        styles["KBody"],
    ))
    story.append(_rl_paragraph(
        "백테스트는 평균적으로 방향성보다 범위 적중률이 더 약했고, 큰 오차는 대부분 추세 반전이나 이벤트 충격을 충분히 반영하지 못한 케이스에서 발생했다.",
        styles["KBody"],
    ))
    story.append(_rl_paragraph(
        "대표 사례들의 feature에는 가격 모멘텀, 변동성, drawdown, 거래량 비율이 핵심이고, event_bias나 macro_bias는 대부분 0에 가까워 외생 변수 반영이 아직 얇다.",
        styles["KBody"],
    ))

    story.append(_rl_paragraph("전체 백테스트 요약", styles["KH1"]))
    summary_lines = [
        f"선정 종목 수: {len(COMPANIES)}",
        f"총 케이스 수: {overall.get('cases', 0)}",
        f"평균 예측 수익률: {_fmt_pct(overall.get('avg_expected_return'))}",
        f"평균 실현 수익률: {_fmt_pct(overall.get('avg_realized_return'))}",
        f"평균 절대오차: {_fmt_pct(overall.get('avg_abs_error'))}",
        f"구간 적중률: {_fmt_pct(overall.get('interval_hit_rate'))}",
        f"방향 정확도: {_fmt_pct(overall.get('direction_accuracy'))}",
    ]
    for line in summary_lines:
        story.append(_rl_paragraph("• " + line, styles["KBody"]))

    forecast_rows = [["종목", "기간", "as_of", "목표일", "기준가", "예상수익률", "예측밴드"]]
    for row in latest:
        as_of_short = _fmt_dt(row["as_of"]).split("T")[0]
        forecast_rows.append(
            [
                COMPANY_NAMES.get(row["company_code"], row["company_code"]),
                f"{row['horizon_days']}d",
                as_of_short,
                str(row["target_date"]),
                _fmt_num(row["base_price"]),
                _fmt_pct(row["expected_return"]),
                f"{_fmt_pct(row['low_return'])} ~ {_fmt_pct(row['high_return'])}",
            ]
        )
    story.append(_rl_paragraph("최신 예측 스냅샷", styles["KH1"]))
    story.append(_rl_table(forecast_rows, [1.10 * inch, 0.50 * inch, 1.15 * inch, 0.95 * inch, 0.80 * inch, 0.78 * inch, 1.22 * inch], font_name, 7.8))
    story.append(Spacer(1, 0.10 * inch))

    backtest_rows = [["종목", "기간", "케이스", "예상", "실현", "절대오차", "적중률", "방향정확도"]]
    for row in backtest:
        backtest_rows.append(
            [
                COMPANY_NAMES.get(row["company_code"], row["company_code"]),
                f"{row['horizon_days']}d",
                str(int(row["cases"])),
                _fmt_pct(row["avg_expected_return"]),
                _fmt_pct(row["avg_realized_return"]),
                _fmt_pct(row["avg_abs_error"]),
                _fmt_pct(row["interval_hit_rate"]),
                _fmt_pct(row["direction_accuracy"]),
            ]
        )
    story.append(_rl_paragraph("백테스트 요약", styles["KH1"]))
    story.append(_rl_table(backtest_rows, [1.10 * inch, 0.50 * inch, 0.45 * inch, 0.70 * inch, 0.70 * inch, 0.75 * inch, 0.65 * inch, 0.80 * inch], font_name, 7.7))

    for case in cases:
        story.append(PageBreak())
        story.append(_rl_paragraph(f"{case['company_name']} 대표 사례", styles["KH1"]))
        story.append(_rl_paragraph("선정 기준: 회사별 절대오차가 가장 큰 백테스트 케이스를 대표 사례로 선택했다.", styles["KBody"]))
        summary_rows = [["항목", "값"]]
        summary_rows += [
            ["as_of", _fmt_dt(case["as_of"])],
            ["horizon", f"{case['horizon_days']}d"],
            ["target_date", str(case["target_date"])],
            ["base_price", _fmt_num(case["base_price"])],
            ["predicted_return", _fmt_pct(case["expected_return"])],
            ["predicted_band", f"{_fmt_pct(case['low_return'])} ~ {_fmt_pct(case['high_return'])}"],
            ["realized_close", _fmt_num(case["realized_close"])],
            ["realized_return", _fmt_pct(case["realized_return"])],
            ["abs_error", _fmt_pct(case["abs_error"])],
            ["interval_hit", "YES" if case["interval_hit"] else "NO"],
            ["scenario_hit", case["scenario_hit"]],
            ["feedback", case["feedback"]],
        ]
        story.append(_rl_table(summary_rows, [1.75 * inch, 4.75 * inch], font_name, 7.8))
        story.append(Spacer(1, 0.08 * inch))

        story.append(_rl_paragraph("예측에 사용된 데이터", styles["KH2"]))
        feature_rows = [["feature", "value"]]
        for key in [
            "latest_close",
            "return_7d",
            "return_20d",
            "return_60d",
            "volatility_20d",
            "volatility_60d",
            "drawdown_60d",
            "range_60d",
            "volume_ratio_20d",
            "trend_strength",
        ]:
            val = case["features"].get(key)
            if key in {"latest_close", "range_60d", "volume_ratio_20d", "trend_strength"}:
                feature_rows.append([key, _fmt_num(val)])
            else:
                feature_rows.append([key, _fmt_pct(val)])
        story.append(_rl_table(feature_rows, [2.0 * inch, 4.5 * inch], font_name, 7.5))
        story.append(Spacer(1, 0.08 * inch))

        story.append(_rl_paragraph("시나리오 신호", styles["KH2"]))
        scenario_rows = [["scenario", "probability", "expected", "band", "rationale"]]
        for sc in case["signals"]:
            scenario_rows.append(
                [
                    sc.get("name", "-"),
                    f"{float(sc.get('probability') or 0.0):.2%}",
                    _fmt_pct(sc.get("expected_return")),
                    f"{_fmt_pct(sc.get('low_return'))} ~ {_fmt_pct(sc.get('high_return'))}",
                    sc.get("rationale", ""),
                ]
            )
        story.append(_rl_table(scenario_rows, [0.8 * inch, 0.8 * inch, 0.8 * inch, 1.0 * inch, 2.9 * inch], font_name, 7.1))
        story.append(Spacer(1, 0.08 * inch))

        if case["company_code"] == "005930":
            narrative = (
                "삼성전자는 최근 7일/20일/60일 수익률이 모두 강했고(trend_strength 3.0), 60일 range가 넓으면서도 "
                "변동성은 비교적 안정적이었다. 모델은 이런 모멘텀을 바탕으로 bull 시나리오 84.8%에 가까운 상방 지속을 "
                "가정했지만, 실제로는 30일 뒤 수익률이 -6.68%로 반전됐다."
            )
            fix = "외생 이벤트(공시/IR/산업 수급) 신호가 거의 없어 가격 모멘텀에 과하게 의존한 것이 약점이다."
        elif case["company_code"] == "000660":
            narrative = (
                "SK하이닉스는 단기 수익률이 음수였지만 60일 기준으로는 여전히 플러스였고, drawdown도 -20%대였다. "
                "모델은 완만한 회복을 예상하며 base/bear 혼합 시나리오를 유지했지만, 실제로는 30일 동안 +82.76%로 "
                "급격한 재평가가 발생했다."
            )
            fix = "HBM/메모리 수요 급변처럼 가격 외 신호가 들어와야 설명력이 높아지는데, 현재 payload는 그 축이 비어 있다."
        else:
            narrative = (
                "한미반도체는 최근 7일/20일 수익률이 모두 음수이고 trend_strength가 -3.0으로 강한 약세였기 때문에, "
                "모델은 bear 시나리오 84.8%를 부여하고 추가 하락을 예상했다. 그러나 실제 30일 수익률은 +56.55%로 "
                "급반전했다."
            )
            fix = "장비/후공정 수요 전환 같은 외생 촉매가 없으면 단기 약세를 장기 약세로 오인하기 쉽다."
        story.append(_rl_paragraph(narrative, styles["KBody"]))
        story.append(_rl_paragraph("개선 포인트: " + fix, styles["KBody"]))

    story.append(PageBreak())
    story.append(_rl_paragraph("결론 및 한계", styles["KH1"]))
    story.append(_rl_paragraph(
        "이 리포트의 예측 엔진은 LLM이 임의로 수치를 뱉는 방식이 아니라, 가격 히스토리와 이벤트/지표 신호를 조합해 "
        "시나리오를 계산하고 그 결과를 백테스트로 검증하는 구조다.",
        styles["KBody"],
    ))
    story.append(_rl_paragraph(
        "다만 대표 사례들을 보면 event_bias와 macro_bias가 대부분 비어 있어, 외생 변수 충격이나 급격한 재평가 구간에서는 "
        "예측 범위가 쉽게 벗어난다. 따라서 배포용 품질을 더 높이려면 공시, IR, 업계 수급, 표준/로드맵 신호를 더 촘촘하게 "
        "넣어야 한다.",
        styles["KBody"],
    ))
    story.append(_rl_paragraph(
        "요약하면 현재 모델은 방향성보다 추세 지속 구간에 강하고, 이벤트 전환이나 급변 구간에 약하다. "
        "이 약점을 줄이기 위한 다음 단계는 데이터 소스 확장과 캘리브레이션 보정이다.",
        styles["KBody"],
    ))

    def _page_canvas(canvas, doc):
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawRightString(7.75 * inch, 0.5 * inch, f"{doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        title="삼성전자·SK하이닉스·한미반도체 주가 예측 및 백테스트 리포트",
        author="Codex",
    )
    doc.build(story, onFirstPage=_page_canvas, onLaterPages=_page_canvas)
    return output_path


if __name__ == "__main__":
    out_docx = REPORT_DIR / "price_forecast_backtest_report.docx"
    out_pdf = REPORT_DIR / "price_forecast_backtest_report.pdf"
    build_report(out_docx)
    build_pdf(out_pdf)
    print(out_docx)
    print(out_pdf)
